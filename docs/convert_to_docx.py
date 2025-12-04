#!/usr/bin/env python3
"""
Convert reliability_report_summary.md to DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def parse_markdown_table(table_text):
    """Parse markdown table into rows and columns"""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    if not lines:
        return []
    
    # Skip separator line (|---|---|)
    data_lines = [line for line in lines if not re.match(r'^[\|\s\-\:]+$', line)]
    
    rows = []
    for line in data_lines:
        # Split by | and clean
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last cells from markdown format
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)
    
    return rows

def add_table_to_doc(doc, table_text):
    """Add a markdown table to the document"""
    rows = parse_markdown_table(table_text)
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Populate table
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # Clear default paragraph
            cell.paragraphs[0].clear()
            paragraph = cell.paragraphs[0]
            
            # Handle bold text in cells
            cell_text = cell_data
            # Find all bold sections
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    # Regular text
                    paragraph.add_run(part)
            
            # Make header row bold
            if i == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    table_buffer = []
    in_table = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Handle horizontal rules
        elif line.strip() == '---':
            if in_table and table_buffer:
                add_table_to_doc(doc, '\n'.join(table_buffer))
                table_buffer = []
                in_table = False
            # Add spacing
            doc.add_paragraph()
        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            in_table = True
            table_buffer.append(line)
        # Handle end of table
        elif in_table and not line.strip().startswith('|'):
            if table_buffer:
                add_table_to_doc(doc, '\n'.join(table_buffer))
                table_buffer = []
            in_table = False
            # Process the current line as regular text
            if line.strip() and not line.startswith('#'):
                p = doc.add_paragraph()
                # Remove markdown formatting
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                text = re.sub(r'`(.*?)`', r'\1', text)
                run = p.add_run(text)
                run.bold = bool(re.search(r'\*\*', line))
        # Handle regular paragraphs
        elif line.strip() and not in_table:
            if not line.strip().startswith('#'):
                p = doc.add_paragraph()
                
                # Handle inline formatting (bold, code, italic)
                text = line
                # Split by markdown formatting
                parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Code/inline code
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        # Italic text
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part:
                        # Regular text
                        p.add_run(part)
        
        i += 1
    
    # Handle table at end of file
    if in_table and table_buffer:
        add_table_to_doc(doc, '\n'.join(table_buffer))
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    convert_markdown_to_docx('reliability_report_summary.md', 'reliability_report_summary1.docx')

